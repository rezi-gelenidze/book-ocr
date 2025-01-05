import os
import config


def save_as_pdf(pages: [str], output_path: str):
    from fpdf import FPDF

    # Initialize PDF
    pdf = FPDF()
    pdf.add_font('Georgian', '', config.FONT_PATH, uni=True)
    pdf.set_font('Georgian', size=12)

    # Add each page as a separate page in the PDF
    for page_text in pages:
        pdf.add_page()
        pdf.multi_cell(0, 10, page_text)

    # Save the output PDF
    pdf.output(output_path)
    print("Extracted text saved to:", output_path)


def save_as_txt(pages: [str], output_path: str):
    with open(output_path, 'w', encoding='utf-8') as f:
        for page_text in pages:
            f.write(page_text + '\n\n')
            # Add delimiter
            f.write("-" * 80 + "\n\n")

    print("Extracted text saved to:", output_path)


def save_as_docx(pages: [str], output_path: str):
    from docx import Document

    doc = Document()
    for page_text in pages:
        cleaned_text = ''.join(c for c in page_text if c.isprintable())  # Clean the text
        doc.add_paragraph(cleaned_text)  # Add page text
        doc.add_paragraph("-" * 80)  # Add delimiter

    doc.save(output_path)
    print("Extracted text saved to:", output_path)


def assemble(pages: [str]):
    # Determine the output type and save the extracted text accordingly

    if config.DEFAULT_OUTPUT_TYPE == config.OUTPUT_TYPES.PDF:
        print("Assembling PDF...")
        save_as_pdf(pages, os.path.join(config.OUTPUT_DIR, 'output.pdf'))
    elif config.DEFAULT_OUTPUT_TYPE == config.OUTPUT_TYPES.TXT:
        save_as_txt(pages, os.path.join(config.OUTPUT_DIR, 'output.txt'))
    elif config.DEFAULT_OUTPUT_TYPE == config.OUTPUT_TYPES.DOCX:
        save_as_docx(pages, os.path.join(config.OUTPUT_DIR, 'output.docx'))
